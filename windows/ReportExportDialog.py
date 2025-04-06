from PySide6.QtWidgets import QDialog, QVBoxLayout, QTextEdit, QPushButton, QFileDialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class ReportExportDialog(QDialog):
    def __init__(self, report_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("报告预览与导出")
        self.setMinimumSize(800, 600)

        # 控件
        self.preview_edit = QTextEdit()
        self.preview_edit.setReadOnly(True)
        self.export_btn = QPushButton("导出为Word")
        self.export_btn.clicked.connect(self.export_to_word)

        # 布局
        layout = QVBoxLayout()
        layout.addWidget(self.preview_edit)
        layout.addWidget(self.export_btn)
        self.setLayout(layout)

        # 生成报告内容
        self.report_data = report_data
        self.generate_preview()

    def generate_preview(self):
        """生成预览文本"""
        content = (
            "====== 肺癌诊断报告 ======\n\n"
            f"患者信息:\n"
            f"姓名: {self.report_data['patient_info']['name']}\n"
            f"性别: {self.report_data['patient_info']['gender']}\n"
            f"年龄: {self.report_data['patient_info']['age']}\n"
            f"病史: {self.report_data['patient_info']['medical_history']}\n\n"
            f"结节特征:\n{self.report_data['nodule_features']}\n\n"
            f"分期预测：\n{self.report_data['stage_prediction']}\n\n"
            f"临床建议:\n{self.report_data['clinical_advice']}"
        )
        self.preview_edit.setPlainText(content)

    def export_to_word(self):
        """导出为Word文档"""
        doc = Document()

        # 添加标题（居中+加粗）
        title = doc.add_paragraph('肺癌诊断报告')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.style.font.size = Pt(16)
        title.runs[0].bold = True

        # 患者信息表格
        patient_info = (
            f"姓名: {self.report_data['patient_info']['name']}\n"
            f"性别: {self.report_data['patient_info']['gender']}\n"
            f"年龄: {self.report_data['patient_info']['age']}\n"
            f"病史: {self.report_data['patient_info']['medical_history']}"
        )
        doc.add_paragraph(patient_info)

        doc.add_paragraph('结节特征:', style="Heading 2").bold = True
        doc.add_paragraph(self.report_data['nodule_features'])
        doc.add_paragraph('分期预测:', style="Heading 2").bold = True
        doc.add_paragraph(self.report_data['stage_prediction'])
        doc.add_paragraph('临床建议:', style="Heading 2").bold = True
        doc.add_paragraph(self.report_data['clinical_advice'])

        # 保存文件
        path, _ = QFileDialog.getSaveFileName(self, "保存报告", self.report_data['patient_info']['name'], "Word 文档 (*.docx)")
        if path:
            doc.save(path)
            self.accept()
